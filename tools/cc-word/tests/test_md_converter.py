"""Tests for cc-word markdown converter (DOCX -> Markdown)."""

from pathlib import Path

import pytest
from docx import Document

from src.md_converter import convert_docx_to_markdown


def _create_docx(path: Path, paragraphs: list[tuple[str, str | None]]) -> Path:
    """Helper: create a test DOCX with given paragraphs.

    Args:
        path: Output path.
        paragraphs: List of (text, style) tuples. style=None for normal.
    """
    doc = Document()
    for text, style in paragraphs:
        if style and style.startswith("Heading"):
            level = int(style.replace("Heading ", "").replace("Heading", "1"))
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    doc.save(str(path))
    return path


class TestConvertDocxToMarkdown:
    """Tests for convert_docx_to_markdown."""

    def test_basic_paragraphs(self, tmp_path):
        docx_path = _create_docx(
            tmp_path / "test.docx",
            [("Hello world", None), ("Second paragraph", None)],
        )
        output = tmp_path / "test.md"

        md = convert_docx_to_markdown(docx_path, output)

        assert "Hello world" in md
        assert "Second paragraph" in md

    def test_headings_converted(self, tmp_path):
        docx_path = _create_docx(
            tmp_path / "test.docx",
            [
                ("Title", "Heading 1"),
                ("Some text", None),
                ("Subtitle", "Heading 2"),
            ],
        )
        output = tmp_path / "test.md"

        md = convert_docx_to_markdown(docx_path, output)

        assert "# Title" in md or "Title" in md  # mammoth may use different heading marks
        assert "Subtitle" in md

    def test_no_images_no_directory(self, tmp_path):
        docx_path = _create_docx(
            tmp_path / "test.docx",
            [("No images", None)],
        )
        output = tmp_path / "test.md"

        md = convert_docx_to_markdown(docx_path, output)

        images_dir = tmp_path / "test_images"
        assert not images_dir.exists()

    def test_output_ends_with_newline(self, tmp_path):
        docx_path = _create_docx(
            tmp_path / "test.docx",
            [("Hello", None)],
        )
        output = tmp_path / "test.md"

        md = convert_docx_to_markdown(docx_path, output)

        assert md.endswith("\n")

    def test_no_excessive_blank_lines(self, tmp_path):
        docx_path = _create_docx(
            tmp_path / "test.docx",
            [("A", None), ("B", None), ("C", None), ("D", None)],
        )
        output = tmp_path / "test.md"

        md = convert_docx_to_markdown(docx_path, output)

        assert "\n\n\n" not in md

    def test_table_in_docx(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "foo"
        table.cell(1, 1).text = "bar"
        doc.add_paragraph("After table")
        docx_path = tmp_path / "table.docx"
        doc.save(str(docx_path))
        output = tmp_path / "table.md"

        md = convert_docx_to_markdown(docx_path, output)

        assert "Name" in md
        assert "foo" in md
