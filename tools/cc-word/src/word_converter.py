"""Word document conversion using python-docx with theme support."""

from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import canonical themes - handle both package and frozen modes
try:
    from cc_shared.themes import CanonicalTheme, get_theme
except ImportError:
    try:
        from themes import CanonicalTheme, get_theme
    except ImportError:
        CanonicalTheme = None
        get_theme = None


def _hex_to_rgb(hex_color: str) -> Optional[RGBColor]:
    """Convert hex color string to RGBColor. Returns None for non-hex values."""
    hex_color = hex_color.strip()
    if not hex_color.startswith("#"):
        return None
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c * 2 for c in hex_color)
    if len(hex_color) != 6:
        return None
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except ValueError:
        return None


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    hex_clean = hex_color.lstrip("#")
    if len(hex_clean) != 6:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_clean)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _set_paragraph_shading(para, hex_color: str) -> None:
    """Apply background shading to a paragraph."""
    hex_clean = hex_color.lstrip("#")
    if len(hex_clean) != 6:
        return
    pPr = para._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_clean)
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)


def _add_paragraph_border_bottom(para, color_hex: str, size: int = 6) -> None:
    """Add a bottom border to a paragraph (used for horizontal rules)."""
    hex_clean = color_hex.lstrip("#")
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_clean)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_paragraph_border_left(para, color_hex: str, size: int = 12) -> None:
    """Add a left border to a paragraph (used for blockquotes)."""
    hex_clean = color_hex.lstrip("#")
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_clean)
    pBdr.append(left)
    pPr.append(pBdr)


def _line_height_to_spacing(line_height_str: str) -> Optional[Pt]:
    """Convert CSS line-height string to Word line spacing value."""
    try:
        val = float(line_height_str)
        # python-docx uses Pt for line spacing when set via line_spacing
        # A value of Pt(X) with line_spacing_rule = EXACTLY uses exact points
        # For proportional, we use the float directly on line_spacing
        return val
    except (ValueError, TypeError):
        return None


def convert_to_word(
    html_content: str,
    output_path: Path,
    theme_name: str = "paper",
) -> None:
    """Convert HTML to Word document with theme styling.

    Maps HTML elements to Word styles:
    - h1-h6 -> Heading 1-6
    - p -> Normal
    - ul/ol -> List styles
    - table -> Table with theme colors
    - pre/code -> Code style with theme code font
    - blockquote -> Quote style
    - hr -> Paragraph with bottom border

    Args:
        html_content: Complete HTML document string
        output_path: Path for output .docx file
        theme_name: Theme to apply for fonts and colors
    """
    # Get theme
    theme = None
    if get_theme is not None:
        try:
            theme = get_theme(theme_name)
        except ValueError:
            pass

    # Parse HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # Create document
    doc = Document()

    # Apply theme fonts and spacing to default style
    if theme:
        style = doc.styles["Normal"]
        style.font.name = theme.fonts.body
        text_color = _hex_to_rgb(theme.colors.text)
        if text_color:
            style.font.color.rgb = text_color

        # Set paragraph spacing from theme line height
        line_height = _line_height_to_spacing(theme.style.body_line_height)
        if line_height:
            style.paragraph_format.line_spacing = line_height
        style.paragraph_format.space_after = Pt(6)

        # Apply heading fonts and spacing
        for level in range(1, 7):
            style_name = f"Heading {level}"
            if style_name in [s.name for s in doc.styles]:
                h_style = doc.styles[style_name]
                h_style.font.name = theme.fonts.heading
                heading_color = _hex_to_rgb(theme.colors.heading)
                if heading_color:
                    h_style.font.color.rgb = heading_color
                h_style.paragraph_format.space_before = Pt(18 - level * 2)
                h_style.paragraph_format.space_after = Pt(6)

    # Find the main content
    body = soup.find("article") or soup.find("body") or soup

    # Process elements
    _process_element(doc, body, theme)

    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Save
    doc.save(str(output_path))


def _process_element(doc: Document, element, theme: Optional[CanonicalTheme] = None):
    """Recursively process HTML elements."""
    if element.name is None:
        return

    for child in element.children:
        if child.name is None:
            if child.string and child.string.strip():
                pass
            continue

        if child.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(child.name[1])
            text = child.get_text(strip=True)
            heading = doc.add_heading(text, level=level)

            # Apply theme accent underline on h1
            if theme and level == 1:
                primary_color = _hex_to_rgb(theme.colors.primary)
                if primary_color and heading.runs:
                    heading.runs[0].font.color.rgb = primary_color

        elif child.name == "p":
            text = child.get_text(strip=True)
            if text:
                doc.add_paragraph(text)

        elif child.name == "ul":
            _process_list(doc, child, ordered=False)

        elif child.name == "ol":
            _process_list(doc, child, ordered=True)

        elif child.name == "blockquote":
            _process_blockquote(doc, child, theme)

        elif child.name == "pre":
            _process_code_block(doc, child, theme)

        elif child.name == "table":
            _process_table(doc, child, theme)

        elif child.name == "hr":
            _process_hr(doc, theme)

        elif child.name in ["div", "article", "section", "main"]:
            _process_element(doc, child, theme)


def _process_list(doc: Document, list_element, ordered: bool = False):
    """Process ul or ol list."""
    for li in list_element.find_all("li", recursive=False):
        text = li.get_text(strip=True)
        if text:
            style = "List Number" if ordered else "List Bullet"
            try:
                doc.add_paragraph(text, style=style)
            except KeyError:
                para = doc.add_paragraph(text)
                para.paragraph_format.left_indent = Inches(0.5)


def _process_blockquote(doc: Document, element, theme: Optional[CanonicalTheme] = None):
    """Process blockquote with left border and optional italic styling."""
    text = element.get_text(strip=True)
    if not text:
        return

    para = doc.add_paragraph(text)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    if theme:
        # Apply left border using theme blockquote border color
        border_color = theme.colors.blockquote_border
        if border_color.startswith("#"):
            _add_paragraph_border_left(para, border_color)

        # Apply blockquote text color
        bq_color = _hex_to_rgb(theme.colors.blockquote_text)
        if bq_color:
            for run in para.runs:
                run.font.color.rgb = bq_color

        # Italic for boardroom and thesis themes
        if theme.name in ("boardroom", "thesis"):
            for run in para.runs:
                run.font.italic = True


def _process_code_block(doc: Document, element, theme: Optional[CanonicalTheme] = None):
    """Process code block with monospace font and background shading."""
    code_text = element.get_text()
    para = doc.add_paragraph()
    run = para.add_run(code_text)

    # Use theme code font
    if theme:
        run.font.name = theme.fonts.code
    else:
        run.font.name = "Consolas"
    run.font.size = Pt(9)

    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    # Apply code background shading
    if theme:
        code_bg = theme.colors.code_bg
        if code_bg.startswith("#"):
            _set_paragraph_shading(para, code_bg)

        code_color = _hex_to_rgb(theme.colors.code_text)
        if code_color:
            run.font.color.rgb = code_color


def _process_hr(doc: Document, theme: Optional[CanonicalTheme] = None):
    """Process horizontal rule as a paragraph with bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    if theme:
        border_color = theme.colors.border
        if theme.name == "boardroom":
            border_color = theme.colors.accent
        if border_color.startswith("#"):
            _add_paragraph_border_bottom(para, border_color, size=6)
    else:
        _add_paragraph_border_bottom(para, "CCCCCC", size=6)


def _process_table(doc: Document, table_element, theme: Optional[CanonicalTheme] = None):
    """Process HTML table with theme colors and alternating rows."""
    rows = table_element.find_all("tr")
    if not rows:
        return

    first_row = rows[0]
    cols = first_row.find_all(["th", "td"])
    num_cols = len(cols)

    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for col_idx, cell in enumerate(cells):
            if col_idx < num_cols:
                text = cell.get_text(strip=True)
                table.rows[row_idx].cells[col_idx].text = text

                if cell.name == "th":
                    for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            # Apply theme header colors
                            if theme:
                                header_text_color = _hex_to_rgb(theme.colors.table_header_text)
                                if header_text_color:
                                    run.font.color.rgb = header_text_color
                                run.font.name = theme.fonts.heading

                    # Apply theme header background
                    if theme:
                        header_bg = theme.colors.table_header_bg
                        if header_bg.startswith("#"):
                            _set_cell_shading(table.rows[row_idx].cells[col_idx], header_bg)

                # Alternating row shading (skip header row)
                elif theme and cell.name == "td" and row_idx % 2 == 0 and row_idx > 0:
                    alt_bg = theme.colors.alt_row_bg
                    if alt_bg.startswith("#"):
                        _set_cell_shading(table.rows[row_idx].cells[col_idx], alt_bg)
